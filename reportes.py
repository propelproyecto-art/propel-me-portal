"""
reportes.py — Genera el borrador de reporte completo a partir de la tabla maestra.

Orquesta el flujo:
  1. Tabla maestra → llamadas al LLM por cada indicador (paralelo)
  2. LLM genera resumen ejecutivo
  3. Ensamblaje en plantilla institucional
  4. Exportación a Markdown / Word / Google Docs
"""
import pandas as pd
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed
import llm_client


# ============================================================
# AGRUPACIÓN DE INDICADORES POR SECCIÓN (replicando estructura del PDF C8)
# ============================================================
SECCIONES = [
    ('1. Sociodemográfico', []),  # Solo metadata, no requiere insights
    ('2. NPS (Net Promoter Score)', ['NPS']),
    ('3. Eficiencia', [
        '% participantes mejoraron eficiencia',
        'Promedio horas ahorradas/semana (calculado pre-post)',
        'Promedio horas ahorradas/semana (proyección endline)',
    ]),
    ('4. Adopción de IA', [
        'AI Adoption Level - Estratégico',
        'AI Adoption Level - Integrado',
        'AI Adoption Level - Activo',
        'AI Adoption Level - Explorando',
        '% orgs aumentaron Net AI Adoption',
        '% orgs nivel Avanzado Net AI Adoption',
    ]),
    ('5. Madurez Digital', [
        '% orgs mejoraron Digital Maturity (total)',
        '% orgs mejoraron DM dimensión D1',
        '% orgs mejoraron DM dimensión D2',
        '% orgs mejoraron DM dimensión D3',
        '% orgs mejoraron DM dimensión D4',
        '% orgs mejoraron DM dimensión D5',
        '% orgs mejoraron DM dimensión D6',
    ]),
    ('6. Uso de herramientas Google AI', [
        '% participantes uso diario Google AI',
    ]),
    ('7. AI Mindset Index', [
        '% participantes con AI Mindset alto',
    ]),
    ('8. Tool Learning', [
        'Tool Learning Marketing - promedio',
        'Tool Learning Marketing - % aprendizaje sig.',
        'Tool Learning Impacto - promedio',
        'Tool Learning Impacto - % aprendizaje sig.',
        'Tool Learning Eficiencia - promedio',
        'Tool Learning Eficiencia - % aprendizaje sig.',
        'Tool Learning Fundraising - promedio',
        'Tool Learning Fundraising - % aprendizaje sig.',
    ]),
    ('9. Community Building', [
        '% participantes mayor confianza tecnología',
        '% participantes establecieron contacto útil',
        '% participantes adoptó nueva herramienta',
    ]),
]

# Indicadores clave para el resumen ejecutivo (los más relevantes)
INDICADORES_CLAVE_RESUMEN = [
    'NPS',
    '% participantes con AI Mindset alto',
    '% participantes mejoraron eficiencia',
    'Promedio horas ahorradas/semana (proyección endline)',
    '% orgs aumentaron Net AI Adoption',
    '% orgs mejoraron Digital Maturity (total)',
    '% participantes uso diario Google AI',
]


# ============================================================
# GENERACIÓN DE INSIGHTS (con paralelismo para velocidad)
# ============================================================
def _generar_un_insight(row, cohorte, programa):
    """Wrapper para usar en paralelo."""
    try:
        insight = llm_client.generar_insight(
            indicador=row['indicador'],
            valor=row['valor'],
            unidad=row['unidad'] or '',
            n=row['n'],
            detalle=row['detalle'] or '',
            cohorte=cohorte,
            programa=programa,
        )
        return row['indicador'], insight, None
    except Exception as e:
        return row['indicador'], None, str(e)


def generar_todos_los_insights(tabla_maestra, cohorte, programa, progress_callback=None):
    """
    Genera un insight para cada indicador en la tabla.
    Usa concurrencia (max 5 a la vez) para no saturar la API.

    Retorna: dict {indicador: insight}
    """
    insights = {}
    errores = {}
    total = len(tabla_maestra)
    completados = 0

    with ThreadPoolExecutor(max_workers=5) as executor:
        futures = {
            executor.submit(_generar_un_insight, row, cohorte, programa): row['indicador']
            for _, row in tabla_maestra.iterrows()
        }
        for future in as_completed(futures):
            indicador, insight, error = future.result()
            if insight:
                insights[indicador] = insight
            else:
                errores[indicador] = error
            completados += 1
            if progress_callback:
                progress_callback(completados, total, indicador)

    return insights, errores


def generar_resumen_ejecutivo(tabla_maestra, cohorte, programa):
    """Genera el resumen ejecutivo del reporte."""
    indicadores_clave = []
    for ind_nombre in INDICADORES_CLAVE_RESUMEN:
        row = tabla_maestra[tabla_maestra['indicador'] == ind_nombre]
        if not row.empty:
            indicadores_clave.append(row.iloc[0].to_dict())

    # n_part = N del NPS o el más grande
    n_part = int(tabla_maestra['n'].max()) if len(tabla_maestra) else 0
    # n_orgs = N de algún indicador a nivel org
    org_inds = tabla_maestra[tabla_maestra['indicador'].str.contains('orgs ', na=False)]
    n_orgs = int(org_inds['n'].iloc[0]) if not org_inds.empty else 0

    return llm_client.generar_resumen_ejecutivo(
        indicadores_clave, cohorte, programa, n_part, n_orgs
    )


# ============================================================
# ENSAMBLAJE DEL REPORTE COMPLETO
# ============================================================
def ensamblar_reporte(tabla_maestra, insights, resumen_ejecutivo, cohorte, programa):
    """
    Construye un dict con la estructura del reporte para ser exportado a
    distintos formatos.
    """
    n_part = int(tabla_maestra['n'].max()) if len(tabla_maestra) else 0
    org_inds = tabla_maestra[tabla_maestra['indicador'].str.contains('orgs ', na=False)]
    n_orgs = int(org_inds['n'].iloc[0]) if not org_inds.empty else 0

    # Datos sintéticos de enriquecimiento (mock de Salesforce).
    # Cuando se conecte Salesforce real, reemplazar esta sección por una
    # llamada al CRM con la lista de orgs del programa+cohorte.
    try:
        from datos_sinteticos import enriquecer_orgs, agregar_metadata, ORGS_C8
        orgs_lista = list(ORGS_C8.keys())
        orgs_enriq = enriquecer_orgs(orgs_lista)
        meta_socio = agregar_metadata(orgs_enriq)
        # Lista de causas y países en formato legible
        causas_str = ', '.join(
            sorted(meta_socio['causas'].keys(),
                   key=lambda c: -meta_socio['causas'][c])
        )
        paises_str = ', '.join(
            sorted(meta_socio['paises'].keys(),
                   key=lambda p: -meta_socio['paises'][p])
        )
        alcance_str = f"{meta_socio['alcance_total']:,} personas"
    except Exception:
        causas_str = '— (pendiente: enriquecer desde Salesforce)'
        paises_str = '— (pendiente: enriquecer desde Salesforce)'
        alcance_str = '— (pendiente: enriquecer desde Salesforce)'

    secciones_output = []

    # Sección sociodemográfica (sin insight, solo metadata)
    secciones_output.append({
        'titulo': '1. Sociodemográfico',
        'contenido': [
            {
                'tipo': 'metadata',
                'datos': {
                    'Participantes': n_part,
                    'Organizaciones': n_orgs,
                    'Causas sociales': causas_str,
                    'Países': paises_str,
                    'Alcance total': alcance_str,
                }
            }
        ],
    })

    # Resto de secciones
    for titulo, indicadores_seccion in SECCIONES[1:]:
        contenido = []
        for ind_nombre in indicadores_seccion:
            row = tabla_maestra[tabla_maestra['indicador'] == ind_nombre]
            if row.empty:
                continue
            r = row.iloc[0]
            contenido.append({
                'tipo': 'indicador',
                'nombre': ind_nombre,
                'valor': r['valor'],
                'unidad': r['unidad'] or '',
                'n': r['n'],
                'detalle': r['detalle'] or '',
                'insight': insights.get(ind_nombre, '(insight no disponible)'),
            })
        if contenido:
            secciones_output.append({'titulo': titulo, 'contenido': contenido})

    return {
        'titulo': f'Reporte de resultados · {programa} · Cohorte {cohorte}',
        'subtitulo': f'Programa: {programa} · Cohorte: {cohorte} · Fecha: {datetime.now().strftime("%d/%m/%Y")}',
        'resumen_ejecutivo': resumen_ejecutivo,
        'secciones': secciones_output,
        'cohorte': cohorte,
        'programa': programa,
    }


# ============================================================
# EXPORTACIÓN A MARKDOWN
# ============================================================
def reporte_a_markdown(reporte):
    """Convierte el reporte a markdown."""
    lines = []
    lines.append(f'# {reporte["titulo"]}\n')
    lines.append(f'*{reporte["subtitulo"]}*\n')
    lines.append('---\n')
    lines.append('## Resumen ejecutivo\n')
    lines.append(reporte['resumen_ejecutivo'] + '\n')
    lines.append('---\n')

    for seccion in reporte['secciones']:
        lines.append(f'## {seccion["titulo"]}\n')
        for item in seccion['contenido']:
            if item['tipo'] == 'metadata':
                for clave, valor in item['datos'].items():
                    lines.append(f'- **{clave}:** {valor}')
                lines.append('')
            else:
                lines.append(f'### {item["nombre"]}')
                lines.append(f'**Valor:** {item["valor"]}{item["unidad"]} · '
                             f'**n:** {item["n"]}')
                if item['detalle']:
                    lines.append(f'*Detalle:* {item["detalle"]}')
                lines.append('')
                lines.append(item['insight'])
                lines.append('')
        lines.append('')

    return '\n'.join(lines)


# ============================================================
# EXPORTACIÓN A WORD (.docx)
# ============================================================
def reporte_a_docx(reporte, output_path):
    """Genera el reporte en formato Word."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Estilos básicos
    estilo = doc.styles['Normal']
    estilo.font.name = 'Calibri'
    estilo.font.size = Pt(11)

    # Título
    titulo = doc.add_heading(reporte['titulo'], level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtítulo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(reporte['subtitulo'])
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Resumen ejecutivo
    doc.add_heading('Resumen ejecutivo', level=1)
    for parr in reporte['resumen_ejecutivo'].split('\n\n'):
        if parr.strip():
            doc.add_paragraph(parr.strip())

    # Secciones
    for seccion in reporte['secciones']:
        doc.add_heading(seccion['titulo'], level=1)
        for item in seccion['contenido']:
            if item['tipo'] == 'metadata':
                for clave, valor in item['datos'].items():
                    p = doc.add_paragraph()
                    run = p.add_run(f'{clave}: ')
                    run.bold = True
                    p.add_run(str(valor))
            else:
                # Encabezado de indicador
                doc.add_heading(item['nombre'], level=2)
                # Línea de valor
                p = doc.add_paragraph()
                run = p.add_run(f'Valor: {item["valor"]}{item["unidad"]}')
                run.bold = True
                p.add_run(f' · n: {item["n"]}')
                # Detalle si hay
                if item['detalle']:
                    p = doc.add_paragraph()
                    run = p.add_run('Detalle: ')
                    run.italic = True
                    p.add_run(item['detalle'])
                # Insight
                doc.add_paragraph(item['insight'])

    doc.save(output_path)
    return output_path
